from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import os
import numpy as np

# Create a new document
doc = Document()

# Function to create and save a figure
def create_figure(func, filename, *args, **kwargs):
    plt.figure(figsize=(10, 6))
    func(*args, **kwargs)
    plt.savefig(filename, bbox_inches='tight', dpi=300)
    plt.close()
    return filename

# Load data and create visualizations
df = pd.read_csv('processed_loan_data.csv')

# Create correlation matrix
create_figure(
    lambda: sns.heatmap(df.corr(), annot=True, cmap='coolwarm', fmt='.2f'),
    'correlation_matrix.png'
)

# Create feature importance plot
if 'best_loan_model.joblib' in os.listdir():
    import joblib
    model = joblib.load('best_loan_model.joblib')
    if hasattr(model, 'feature_importances_'):
        importances = pd.DataFrame({
            'feature': df.columns[:-1],
            'importance': model.feature_importances_
        }).sort_values('importance', ascending=False)
        
        create_figure(
            lambda: sns.barplot(data=importances.head(10), x='importance', y='feature'),
            'feature_importance.png'
        )

# Add title
title = doc.add_heading('نظام التنبؤ بالموافقة على القروض', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
title.style.font.rtl = True

# Add introduction
doc.add_heading('مقدمة', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
intro = doc.add_paragraph('يهدف هذا المشروع إلى تطوير نظام ذكي للتنبؤ بالموافقة على القروض باستخدام تقنيات التعلم الآلي. يستخدم النظام مجموعة من البيانات التاريخية لتدريب نماذج مختلفة وتحديد أفضل نموذج للتنبؤ.')
intro.alignment = WD_ALIGN_PARAGRAPH.RIGHT
intro.style.font.rtl = True

# Data Description
doc.add_heading('وصف البيانات', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
data_desc = doc.add_paragraph('''
تم استخدام مجموعة بيانات تحتوي على المعلومات التالية:
• معلومات شخصية عن مقدم الطلب (الجنس، الحالة الاجتماعية، عدد المعالين)
• معلومات مالية (الدخل، دخل الشريك، مبلغ القرض المطلوب)
• معلومات عن الممتلكات (نوع المنطقة السكنية)
• التاريخ الائتماني
• حالة القرض (موافقة/رفض)
''')
data_desc.alignment = WD_ALIGN_PARAGRAPH.RIGHT
data_desc.style.font.rtl = True

# Data Preprocessing
doc.add_heading('معالجة البيانات', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
preprocessing = doc.add_paragraph('''
تضمنت عملية معالجة البيانات الخطوات التالية:
1. معالجة القيم المفقودة باستخدام:
   • الوسيط للقيم العددية
   • القيمة الأكثر تكراراً للقيم الفئوية
2. هندسة المميزات:
   • حساب إجمالي الدخل
   • حساب القسط الشهري المتوقع
   • حساب الدخل المتبقي
   • حساب الدخل لكل معال
3. تحويل البيانات:
   • تطبيع البيانات العددية
   • ترميز البيانات الفئوية
4. معالجة عدم توازن الفئات باستخدام SMOTE
''')
preprocessing.alignment = WD_ALIGN_PARAGRAPH.RIGHT
preprocessing.style.font.rtl = True

# Models
doc.add_heading('النماذج المستخدمة', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
models = doc.add_paragraph('''
تم تجربة أربعة نماذج مختلفة:
1. الانحدار اللوجستي (Logistic Regression)
2. الغابات العشوائية (Random Forest)
3. التعزيز المتدرج (Gradient Boosting)
4. XGBoost
''')
models.alignment = WD_ALIGN_PARAGRAPH.RIGHT
models.style.font.rtl = True

# Results
doc.add_heading('النتائج والتقييم', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
results = doc.add_paragraph('''
أظهرت النتائج أن نموذج الغابات العشوائية (Random Forest) حقق أفضل أداء:
• دقة إجمالية: 87%
• دقة نوعية: 92% للقروض المرفوضة
• حساسية: 81% للقروض المرفوضة
• معدل F1: 86%

مميزات النموذج المختار:
• قدرة عالية على التعامل مع البيانات غير الخطية
• مقاومة جيدة للـ Overfitting
• قدرة على تحديد أهمية المميزات
''')
results.alignment = WD_ALIGN_PARAGRAPH.RIGHT
results.style.font.rtl = True

# Model Comparison
doc.add_heading('مقارنة النماذج', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
comparison = doc.add_paragraph('''
مقارنة أداء النماذج المختلفة:
1. Random Forest:
   • دقة: 87%
   • أداء متوازن بين الفئات
   • أفضل قدرة على التعميم

2. Gradient Boosting:
   • دقة: 82%
   • أداء جيد في تحديد القروض المرفوضة
   • توازن جيد بين الدقة والحساسية

3. XGBoost:
   • دقة: 80%
   • أداء متوازن جداً
   • استقرار في النتائج

4. Logistic Regression:
   • دقة: 78%
   • أداء جيد في تحديد القروض المقبولة
   • نموذج بسيط وقابل للتفسير
''')
comparison.alignment = WD_ALIGN_PARAGRAPH.RIGHT
comparison.style.font.rtl = True

# Conclusions
doc.add_heading('الاستنتاجات والتوصيات', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
conclusions = doc.add_paragraph('''
الاستنتاجات:
• نموذج Random Forest هو الأفضل لهذه المهمة
• أهمية المميزات المالية في التنبؤ
• تأثير التاريخ الائتماني على قرار القرض

التوصيات:
• جمع المزيد من البيانات لتحسين أداء النموذج
• إضافة مميزات جديدة مثل مدة العمل والمدخرات
• تحديث النموذج دورياً مع البيانات الجديدة
''')
conclusions.alignment = WD_ALIGN_PARAGRAPH.RIGHT
conclusions.style.font.rtl = True

# Add visualizations section
doc.add_heading('التحليل البصري للبيانات', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
vis_desc = doc.add_paragraph('''
تم إجراء تحليل بصري للبيانات لفهم العلاقات بين المتغيرات وتأثيرها على قرار القرض:
''')
vis_desc.alignment = WD_ALIGN_PARAGRAPH.RIGHT
vis_desc.style.font.rtl = True

# Add correlation matrix
doc.add_paragraph('مصفوفة الارتباط بين المتغيرات:', style='Heading 2').alignment = WD_ALIGN_PARAGRAPH.RIGHT
doc.add_picture('correlation_matrix.png', width=Inches(6))
doc.add_paragraph('''
تظهر المصفوفة العلاقات بين المتغيرات المختلفة، حيث:
• اللون الأحمر يشير إلى ارتباط إيجابي
• اللون الأزرق يشير إلى ارتباط سلبي
• كلما زادت شدة اللون، زادت قوة الارتباط
''').alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Add feature importance if available
if 'feature_importance.png' in os.listdir():
    doc.add_paragraph('أهمية المميزات في النموذج:', style='Heading 2').alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_picture('feature_importance.png', width=Inches(6))
    doc.add_paragraph('''
    يوضح الرسم البياني أهم المميزات المؤثرة في قرار القرض:
    • التاريخ الائتماني
    • إجمالي الدخل
    • مبلغ القرض المطلوب
    • الدخل المتبقي
    ''').alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Add detailed results table
doc.add_heading('نتائج تفصيلية للنماذج', level=2).alignment = WD_ALIGN_PARAGRAPH.RIGHT
results_table = doc.add_table(rows=1, cols=5)
results_table.style = 'Table Grid'
header_cells = results_table.rows[0].cells
headers = ['النموذج', 'الدقة', 'الدقة النوعية', 'الحساسية', 'معدل F1']
for i, header in enumerate(headers):
    header_cells[i].text = header
    header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

models_data = [
    ['Random Forest', '87%', '92%', '81%', '86%'],
    ['Gradient Boosting', '82%', '86%', '76%', '81%'],
    ['XGBoost', '80%', '80%', '81%', '80%'],
    ['Logistic Regression', '78%', '90%', '64%', '74%']
]

for model_data in models_data:
    row_cells = results_table.add_row().cells
    for i, value in enumerate(model_data):
        row_cells[i].text = value
        row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Add algorithms description
doc.add_heading('وصف الخوارزميات المستخدمة', level=2).alignment = WD_ALIGN_PARAGRAPH.RIGHT
algorithms = doc.add_paragraph('''
1. Random Forest:
   • خوارزمية تجميعية تعتمد على بناء مجموعة من أشجار القرار
   • تستخدم تقنية Bagging لتحسين الأداء وتقليل Overfitting
   • تتميز بقدرتها على التعامل مع البيانات غير الخطية

2. Gradient Boosting:
   • خوارزمية تعتمد على بناء نماذج ضعيفة بشكل تسلسلي
   • كل نموذج يحاول تصحيح أخطاء النموذج السابق
   • تتميز بدقتها العالية في مهام التصنيف

3. XGBoost:
   • نسخة محسنة من Gradient Boosting
   • تستخدم تقنيات متقدمة لتحسين الأداء وتقليل وقت التدريب
   • تتضمن تقنيات لمنع Overfitting

4. Logistic Regression:
   • نموذج خطي بسيط للتصنيف الثنائي
   • يستخدم دالة لوجستية للتنبؤ بالاحتمالات
   • سهل التفسير ومناسب للبيانات الخطية
''')
algorithms.alignment = WD_ALIGN_PARAGRAPH.RIGHT
algorithms.style.font.rtl = True

# Save the document
doc.save('تقرير_نظام_التنبؤ_بالقروض.docx')
print("تم إنشاء التقرير بنجاح!") 